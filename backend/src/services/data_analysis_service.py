"""
Data Analysis Service for Agent CEO system
Handles data parsing, analysis, and insights generation using LangChain
"""

import os
import json
import logging
from typing import Dict, List, Optional, Any, Union
from datetime import datetime
import pandas as pd
import numpy as np
from io import StringIO, BytesIO
import base64

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import ChatOpenAI
from langchain_anthropic import ChatAnthropic
from langchain.chains.summarize import load_summarize_chain
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate

# Document processing imports
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import openpyxl
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

class DataAnalysisService:
    """Service for data parsing, analysis, and insights generation"""
    
    def __init__(self):
        # Initialize LLM models
        self.openai_model = ChatOpenAI(
            model="gpt-4-turbo",
            temperature=0.3,
            openai_api_key=os.getenv('OPENAI_API_KEY')
        )
        
        self.anthropic_model = ChatAnthropic(
            model="claude-3-sonnet-20240229",
            temperature=0.3,
            anthropic_api_key=os.getenv('ANTHROPIC_API_KEY')
        )
        
        # Text splitter for large documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=4000,
            chunk_overlap=200,
            length_function=len
        )
        
        # Analysis templates
        self.analysis_templates = {
            'data_summary': """
            Analyze the following data and provide a comprehensive summary:
            
            Data: {data}
            Data Type: {data_type}
            Context: {context}
            
            Provide:
            1. Key insights and patterns
            2. Statistical summary (if applicable)
            3. Notable trends or anomalies
            4. Business implications
            5. Recommendations for action
            
            Focus on actionable insights that can drive business decisions.
            """,
            
            'competitive_analysis': """
            Analyze the following competitive data:
            
            Data: {data}
            Company Focus: {company_focus}
            Analysis Scope: {scope}
            
            Provide:
            1. Competitive positioning analysis
            2. Market share insights
            3. Strengths and weaknesses comparison
            4. Opportunity identification
            5. Strategic recommendations
            
            Focus on competitive advantages and market opportunities.
            """,
            
            'financial_analysis': """
            Analyze the following financial data:
            
            Data: {data}
            Time Period: {time_period}
            Business Context: {context}
            
            Provide:
            1. Financial performance summary
            2. Key metrics analysis
            3. Trend identification
            4. Risk assessment
            5. Growth opportunities
            6. Recommendations for improvement
            
            Focus on financial health and growth potential.
            """,
            
            'customer_analysis': """
            Analyze the following customer data:
            
            Data: {data}
            Analysis Focus: {focus}
            Business Goals: {goals}
            
            Provide:
            1. Customer segmentation insights
            2. Behavior pattern analysis
            3. Satisfaction and retention metrics
            4. Growth opportunities
            5. Personalization recommendations
            6. Customer lifetime value insights
            
            Focus on customer experience and revenue optimization.
            """,
            
            'market_analysis': """
            Analyze the following market data:
            
            Data: {data}
            Market Scope: {scope}
            Industry Context: {industry}
            
            Provide:
            1. Market size and growth analysis
            2. Trend identification
            3. Opportunity assessment
            4. Threat analysis
            5. Market entry strategies
            6. Competitive landscape overview
            
            Focus on market opportunities and strategic positioning.
            """
        }
    
    def parse_csv_data(self, csv_content: str, analysis_context: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Parse and analyze CSV data
        
        Args:
            csv_content: CSV data as string
            analysis_context: Context for analysis
            
        Returns:
            Dictionary with parsed data and analysis
        """
        try:
            # Parse CSV
            df = pd.read_csv(StringIO(csv_content))
            
            # Basic statistics
            stats = {
                'row_count': len(df),
                'column_count': len(df.columns),
                'columns': list(df.columns),
                'data_types': df.dtypes.to_dict(),
                'missing_values': df.isnull().sum().to_dict(),
                'numeric_summary': df.describe().to_dict() if len(df.select_dtypes(include=[np.number]).columns) > 0 else {}
            }
            
            # Convert DataFrame to JSON for analysis
            data_sample = df.head(10).to_json(orient='records', indent=2)
            
            # Generate AI analysis
            context = analysis_context or {}
            analysis_result = self._generate_data_analysis(
                data=data_sample,
                data_type='csv',
                context=context,
                stats=stats
            )
            
            return {
                'success': True,
                'data_type': 'csv',
                'statistics': stats,
                'sample_data': json.loads(data_sample),
                'analysis': analysis_result,
                'parsed_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"CSV parsing error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def parse_excel_data(self, file_path: str, sheet_name: str = None, 
                        analysis_context: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Parse and analyze Excel data
        
        Args:
            file_path: Path to Excel file
            sheet_name: Specific sheet to analyze (optional)
            analysis_context: Context for analysis
            
        Returns:
            Dictionary with parsed data and analysis
        """
        try:
            # Read Excel file
            if sheet_name:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_analyzed = [sheet_name]
            else:
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                sheets_analyzed = excel_file.sheet_names
                df = pd.read_excel(file_path, sheet_name=sheets_analyzed[0])  # Analyze first sheet
            
            # Basic statistics
            stats = {
                'sheets_available': sheets_analyzed,
                'current_sheet': sheet_name or sheets_analyzed[0],
                'row_count': len(df),
                'column_count': len(df.columns),
                'columns': list(df.columns),
                'data_types': df.dtypes.to_dict(),
                'missing_values': df.isnull().sum().to_dict(),
                'numeric_summary': df.describe().to_dict() if len(df.select_dtypes(include=[np.number]).columns) > 0 else {}
            }
            
            # Convert DataFrame to JSON for analysis
            data_sample = df.head(10).to_json(orient='records', indent=2)
            
            # Generate AI analysis
            context = analysis_context or {}
            analysis_result = self._generate_data_analysis(
                data=data_sample,
                data_type='excel',
                context=context,
                stats=stats
            )
            
            return {
                'success': True,
                'data_type': 'excel',
                'statistics': stats,
                'sample_data': json.loads(data_sample),
                'analysis': analysis_result,
                'parsed_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Excel parsing error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def parse_json_data(self, json_content: Union[str, dict], 
                       analysis_context: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Parse and analyze JSON data
        
        Args:
            json_content: JSON data as string or dict
            analysis_context: Context for analysis
            
        Returns:
            Dictionary with parsed data and analysis
        """
        try:
            # Parse JSON if string
            if isinstance(json_content, str):
                data = json.loads(json_content)
            else:
                data = json_content
            
            # Analyze structure
            def analyze_structure(obj, path="root"):
                if isinstance(obj, dict):
                    return {
                        'type': 'object',
                        'keys': list(obj.keys()),
                        'key_count': len(obj.keys()),
                        'nested_structure': {k: analyze_structure(v, f"{path}.{k}") for k, v in obj.items() if isinstance(v, (dict, list))}
                    }
                elif isinstance(obj, list):
                    return {
                        'type': 'array',
                        'length': len(obj),
                        'item_types': list(set(type(item).__name__ for item in obj)),
                        'sample_structure': analyze_structure(obj[0], f"{path}[0]") if obj else None
                    }
                else:
                    return {'type': type(obj).__name__, 'value': str(obj)[:100]}
            
            structure_analysis = analyze_structure(data)
            
            # Convert to string for AI analysis
            data_str = json.dumps(data, indent=2)[:5000]  # Limit size for analysis
            
            # Generate AI analysis
            context = analysis_context or {}
            analysis_result = self._generate_data_analysis(
                data=data_str,
                data_type='json',
                context=context,
                stats={'structure': structure_analysis}
            )
            
            return {
                'success': True,
                'data_type': 'json',
                'structure_analysis': structure_analysis,
                'sample_data': data if len(data_str) < 1000 else "Data too large for full display",
                'analysis': analysis_result,
                'parsed_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"JSON parsing error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def parse_pdf_document(self, file_path: str, 
                          analysis_context: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Parse and analyze PDF document
        
        Args:
            file_path: Path to PDF file
            analysis_context: Context for analysis
            
        Returns:
            Dictionary with parsed content and analysis
        """
        try:
            # Extract text from PDF
            reader = PdfReader(file_path)
            text_content = ""
            
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
            
            # Split text into chunks for analysis
            documents = self.text_splitter.split_text(text_content)
            
            # Basic statistics
            stats = {
                'page_count': len(reader.pages),
                'total_characters': len(text_content),
                'total_words': len(text_content.split()),
                'chunks_created': len(documents)
            }
            
            # Generate summary and analysis
            context = analysis_context or {}
            analysis_result = self._generate_document_analysis(
                content=text_content[:5000],  # First 5000 chars for analysis
                document_type='pdf',
                context=context,
                stats=stats
            )
            
            return {
                'success': True,
                'document_type': 'pdf',
                'statistics': stats,
                'content_preview': text_content[:1000],
                'analysis': analysis_result,
                'parsed_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"PDF parsing error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def parse_word_document(self, file_path: str, 
                           analysis_context: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Parse and analyze Word document
        
        Args:
            file_path: Path to Word document
            analysis_context: Context for analysis
            
        Returns:
            Dictionary with parsed content and analysis
        """
        try:
            # Extract text from Word document
            doc = DocxDocument(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Basic statistics
            stats = {
                'paragraph_count': len(doc.paragraphs),
                'total_characters': len(text_content),
                'total_words': len(text_content.split())
            }
            
            # Generate analysis
            context = analysis_context or {}
            analysis_result = self._generate_document_analysis(
                content=text_content[:5000],
                document_type='word',
                context=context,
                stats=stats
            )
            
            return {
                'success': True,
                'document_type': 'word',
                'statistics': stats,
                'content_preview': text_content[:1000],
                'analysis': analysis_result,
                'parsed_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Word document parsing error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _generate_data_analysis(self, data: str, data_type: str, 
                               context: Dict[str, Any], stats: Dict[str, Any]) -> str:
        """Generate AI-powered data analysis"""
        try:
            prompt = self.analysis_templates['data_summary'].format(
                data=data,
                data_type=data_type,
                context=json.dumps(context, indent=2)
            )
            
            # Use OpenAI for analysis
            response = self.openai_model.invoke(prompt)
            return response.content
            
        except Exception as e:
            logger.error(f"Data analysis generation error: {str(e)}")
            return f"Analysis generation failed: {str(e)}"
    
    def _generate_document_analysis(self, content: str, document_type: str,
                                   context: Dict[str, Any], stats: Dict[str, Any]) -> str:
        """Generate AI-powered document analysis"""
        try:
            prompt = f"""
            Analyze the following document content and provide insights:
            
            Document Type: {document_type}
            Content: {content}
            Context: {json.dumps(context, indent=2)}
            Statistics: {json.dumps(stats, indent=2)}
            
            Provide:
            1. Document summary and key themes
            2. Important insights and findings
            3. Business implications
            4. Actionable recommendations
            5. Key data points or metrics mentioned
            
            Focus on extracting business value from the document content.
            """
            
            response = self.openai_model.invoke(prompt)
            return response.content
            
        except Exception as e:
            logger.error(f"Document analysis generation error: {str(e)}")
            return f"Analysis generation failed: {str(e)}"
    
    def generate_competitive_analysis(self, competitor_data: List[Dict[str, Any]], 
                                    company_focus: str) -> Dict[str, Any]:
        """
        Generate competitive analysis from competitor data
        
        Args:
            competitor_data: List of competitor information
            company_focus: Focus area for analysis
            
        Returns:
            Dictionary with competitive analysis
        """
        try:
            data_str = json.dumps(competitor_data, indent=2)
            
            prompt = self.analysis_templates['competitive_analysis'].format(
                data=data_str,
                company_focus=company_focus,
                scope="comprehensive competitive landscape analysis"
            )
            
            response = self.openai_model.invoke(prompt)
            
            return {
                'success': True,
                'analysis_type': 'competitive_analysis',
                'analysis': response.content,
                'competitors_analyzed': len(competitor_data),
                'company_focus': company_focus,
                'generated_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Competitive analysis error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def generate_financial_analysis(self, financial_data: Dict[str, Any], 
                                   time_period: str, context: str) -> Dict[str, Any]:
        """
        Generate financial analysis from financial data
        
        Args:
            financial_data: Financial data dictionary
            time_period: Time period for analysis
            context: Business context
            
        Returns:
            Dictionary with financial analysis
        """
        try:
            data_str = json.dumps(financial_data, indent=2)
            
            prompt = self.analysis_templates['financial_analysis'].format(
                data=data_str,
                time_period=time_period,
                context=context
            )
            
            response = self.openai_model.invoke(prompt)
            
            return {
                'success': True,
                'analysis_type': 'financial_analysis',
                'analysis': response.content,
                'time_period': time_period,
                'context': context,
                'generated_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Financial analysis error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def generate_customer_analysis(self, customer_data: Dict[str, Any], 
                                  focus: str, goals: str) -> Dict[str, Any]:
        """
        Generate customer analysis from customer data
        
        Args:
            customer_data: Customer data dictionary
            focus: Analysis focus area
            goals: Business goals
            
        Returns:
            Dictionary with customer analysis
        """
        try:
            data_str = json.dumps(customer_data, indent=2)
            
            prompt = self.analysis_templates['customer_analysis'].format(
                data=data_str,
                focus=focus,
                goals=goals
            )
            
            response = self.openai_model.invoke(prompt)
            
            return {
                'success': True,
                'analysis_type': 'customer_analysis',
                'analysis': response.content,
                'focus': focus,
                'goals': goals,
                'generated_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Customer analysis error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def generate_market_analysis(self, market_data: Dict[str, Any], 
                                scope: str, industry: str) -> Dict[str, Any]:
        """
        Generate market analysis from market data
        
        Args:
            market_data: Market data dictionary
            scope: Analysis scope
            industry: Industry context
            
        Returns:
            Dictionary with market analysis
        """
        try:
            data_str = json.dumps(market_data, indent=2)
            
            prompt = self.analysis_templates['market_analysis'].format(
                data=data_str,
                scope=scope,
                industry=industry
            )
            
            response = self.openai_model.invoke(prompt)
            
            return {
                'success': True,
                'analysis_type': 'market_analysis',
                'analysis': response.content,
                'scope': scope,
                'industry': industry,
                'generated_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Market analysis error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def health_check(self) -> Dict[str, Any]:
        """
        Check data analysis service health
        
        Returns:
            Dictionary with health status
        """
        health_status = {
            'service_status': 'healthy',
            'openai_configured': bool(os.getenv('OPENAI_API_KEY')),
            'anthropic_configured': bool(os.getenv('ANTHROPIC_API_KEY')),
            'langchain_available': True,
            'supported_formats': ['csv', 'excel', 'json', 'pdf', 'word'],
            'timestamp': datetime.utcnow().isoformat()
        }
        
        # Test LLM connectivity
        try:
            test_response = self.openai_model.invoke("Test message")
            health_status['openai_connection'] = 'healthy'
        except Exception as e:
            health_status['openai_connection'] = f'error: {str(e)}'
        
        return health_status

# Global data analysis service instance
data_analysis_service = DataAnalysisService()

